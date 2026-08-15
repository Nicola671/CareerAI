"""
Document Processor - Extracts text from PDF, DOCX, TXT, and IMAGES (via Groq Vision).
Supports scanned PDFs and photos of documents.
"""
import os
import base64
from typing import List


class DocumentProcessor:
    """Process various document formats and extract text for RAG indexing."""

    SUPPORTED_FORMATS = [".pdf", ".txt", ".docx", ".doc", ".jpg", ".jpeg", ".png", ".webp"]
    IMAGE_FORMATS = [".jpg", ".jpeg", ".png", ".webp", ".gif", ".bmp"]

    @staticmethod
    def extract_text(file_path: str, groq_api_key: str = None) -> str:
        """Extract text from a file based on its extension.
        For images and scanned PDFs, uses Groq Vision API.
        """
        ext = os.path.splitext(file_path)[1].lower()

        if ext in DocumentProcessor.IMAGE_FORMATS:
            if not groq_api_key:
                raise ValueError("Se necesita API key de Groq para procesar imágenes")
            return DocumentProcessor._extract_image(file_path, groq_api_key)
        elif ext == ".pdf":
            return DocumentProcessor._extract_pdf(file_path, groq_api_key)
        elif ext == ".txt":
            return DocumentProcessor._extract_txt(file_path)
        elif ext in [".docx", ".doc"]:
            return DocumentProcessor._extract_docx(file_path)
        else:
            raise ValueError(f"Formato no soportado: {ext}")

    @staticmethod
    def _extract_image(file_path: str, groq_api_key: str) -> str:
        """Extract text from an image using Groq Vision (Llama 4 Scout)."""
        try:
            from groq import Groq

            # Read and encode image
            with open(file_path, "rb") as f:
                image_data = f.read()

            base64_image = base64.b64encode(image_data).decode("utf-8")

            # Detect MIME type
            ext = os.path.splitext(file_path)[1].lower()
            mime_map = {
                ".jpg": "image/jpeg",
                ".jpeg": "image/jpeg",
                ".png": "image/png",
                ".webp": "image/webp",
                ".gif": "image/gif",
                ".bmp": "image/bmp",
            }
            mime_type = mime_map.get(ext, "image/jpeg")

            # Call Groq Vision API
            client = Groq(api_key=groq_api_key)
            response = client.chat.completions.create(
                model="meta-llama/llama-4-scout-17b-16e-instruct",
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": (
                                    "Extraé TODO el texto de esta imagen de documento exactamente como aparece. "
                                    "Incluí todos los detalles: nombres, fechas, experiencia laboral, educación, "
                                    "habilidades, idiomas, certificaciones, datos de contacto, y cualquier otra "
                                    "información. Mantené la estructura original. Si hay tablas, extraé el contenido. "
                                    "Respondé SOLO con el texto extraído, sin comentarios adicionales."
                                ),
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:{mime_type};base64,{base64_image}"
                                },
                            },
                        ],
                    }
                ],
                max_tokens=4096,
                temperature=0.1,
            )

            text = response.choices[0].message.content
            if text and text.strip():
                return text.strip()
            else:
                raise ValueError("No se pudo extraer texto de la imagen")

        except ImportError:
            raise ValueError("Instala el paquete 'groq': pip install groq")
        except Exception as e:
            if "groq" in str(type(e).__module__).lower():
                raise ValueError(f"Error de Groq Vision API: {e}")
            raise ValueError(f"Error procesando imagen: {e}")

    @staticmethod
    def _extract_pdf(file_path: str, groq_api_key: str = None) -> str:
        """Extract text from PDF. Tries 3 methods + Vision API for scanned PDFs."""
        import logging
        logger = logging.getLogger("careerai.document")
        text = ""
        errors = []

        # Method 1: PyPDF (fast, works with text PDFs)
        try:
            from pypdf import PdfReader

            reader = PdfReader(file_path)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            if text.strip() and len(text.strip()) > 20:
                logger.info(f"PDF extracted with pypdf: {len(text)} chars")
                return text.strip()
            else:
                errors.append(f"pypdf: solo {len(text.strip())} chars")
        except ImportError:
            errors.append("pypdf: no instalado")
        except Exception as e:
            errors.append(f"pypdf: {e}")

        # Method 2: pdfplumber (better with complex layouts)
        try:
            import pdfplumber

            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"

                    # Also try extracting tables
                    try:
                        tables = page.extract_tables()
                        for table in tables:
                            for row in table:
                                if row:
                                    row_text = " | ".join(
                                        str(cell).strip() for cell in row if cell
                                    )
                                    if row_text:
                                        text += row_text + "\n"
                    except Exception:
                        pass

            if text.strip() and len(text.strip()) > 20:
                logger.info(f"PDF extracted with pdfplumber: {len(text)} chars")
                return text.strip()
            else:
                errors.append(f"pdfplumber: solo {len(text.strip())} chars")
        except ImportError:
            errors.append("pdfplumber: no instalado")
        except Exception as e:
            errors.append(f"pdfplumber: {e}")

        # Method 3: PyMuPDF / fitz (handles more PDF types)
        try:
            import fitz

            doc = fitz.open(file_path)
            fitz_text = ""
            for page in doc:
                page_text = page.get_text()
                if page_text:
                    fitz_text += page_text + "\n"
            doc.close()

            if fitz_text.strip() and len(fitz_text.strip()) > 20:
                logger.info(f"PDF extracted with fitz: {len(fitz_text)} chars")
                return fitz_text.strip()
            else:
                errors.append(f"fitz/PyMuPDF: solo {len(fitz_text.strip())} chars")
        except ImportError:
            errors.append("PyMuPDF: no instalado")
        except Exception as e:
            errors.append(f"fitz/PyMuPDF: {e}")

        # Method 4: Vision AI - render PDF pages as images and read with Llama Vision
        if groq_api_key:
            try:
                result = DocumentProcessor._extract_pdf_via_vision(
                    file_path, groq_api_key
                )
                logger.info(f"PDF extracted with Vision API: {len(result)} chars")
                return result
            except Exception as vision_err:
                errors.append(f"Vision API: {vision_err}")
        else:
            errors.append("Vision API: no API key disponible")

        # Last resort — return whatever we have
        if text.strip():
            logger.warning(f"PDF fallback con texto parcial: {len(text.strip())} chars")
            return text.strip()

        # Log all errors for debugging
        error_detail = " | ".join(errors)
        logger.error(f"PDF extraction failed for {file_path}: {error_detail}")

        raise ValueError(
            f"No se pudo extraer texto del PDF. "
            f"Intentá subir una imagen/captura del documento. "
            f"(Detalle: {error_detail})"
        )

    @staticmethod
    def _extract_pdf_via_vision(file_path: str, groq_api_key: str) -> str:
        """Extract text from a scanned PDF by converting pages to images and using Vision."""
        import logging
        logger = logging.getLogger("careerai.document")

        VISION_MODELS = [
            "meta-llama/llama-4-scout-17b-16e-instruct",
            "llama-3.2-90b-vision-preview",
            "llama-3.2-11b-vision-preview",
        ]

        # Step 1: Try converting PDF pages to images with PyMuPDF
        try:
            import fitz  # PyMuPDF

            doc = fitz.open(file_path)
            all_text = []

            for page_num in range(min(len(doc), 5)):  # Max 5 pages
                page = doc[page_num]
                # Render page as image
                mat = fitz.Matrix(2, 2)  # 2x zoom for better quality
                pix = page.get_pixmap(matrix=mat)
                img_bytes = pix.tobytes("png")
                base64_image = base64.b64encode(img_bytes).decode("utf-8")

                # Try each vision model
                page_extracted = False
                for model in VISION_MODELS:
                    try:
                        from groq import Groq
                        client = Groq(api_key=groq_api_key)
                        response = client.chat.completions.create(
                            model=model,
                            messages=[
                                {
                                    "role": "user",
                                    "content": [
                                        {
                                            "type": "text",
                                            "text": (
                                                f"Página {page_num + 1}. Extraé TODO el texto de esta página "
                                                "exactamente como aparece. Incluí todos los detalles. "
                                                "Respondé SOLO con el texto extraído."
                                            ),
                                        },
                                        {
                                            "type": "image_url",
                                            "image_url": {
                                                "url": f"data:image/png;base64,{base64_image}"
                                            },
                                        },
                                    ],
                                }
                            ],
                            max_tokens=4096,
                            temperature=0.1,
                        )
                        page_text = response.choices[0].message.content
                        if page_text and page_text.strip():
                            all_text.append(page_text.strip())
                            page_extracted = True
                            break  # Success with this model
                    except Exception as e:
                        logger.warning(f"Vision model {model} failed for page {page_num+1}: {e}")
                        continue

                if not page_extracted:
                    logger.warning(f"All vision models failed for page {page_num+1}")

            doc.close()

            if all_text:
                return "\n\n".join(all_text)

        except ImportError:
            logger.warning("PyMuPDF not installed, skipping page-to-image conversion")
        except Exception as e:
            logger.error(f"PyMuPDF conversion error: {e}")

        raise ValueError("No se pudo extraer texto del PDF con Vision AI")

    @staticmethod
    def _extract_txt(file_path: str) -> str:
        """Extract text from a plain text file."""
        encodings = ["utf-8", "latin-1", "cp1252"]
        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    return f.read().strip()
            except (UnicodeDecodeError, UnicodeError):
                continue
        raise ValueError("No se pudo leer el archivo de texto")

    @staticmethod
    def _extract_docx(file_path: str) -> str:
        """Extract text from a Word document."""
        try:
            from docx import Document

            doc = Document(file_path)
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        paragraphs.append(row_text)

            return "\n".join(paragraphs)
        except Exception as e:
            raise ValueError(f"No se pudo leer el archivo DOCX: {e}")

    @staticmethod
    def chunk_text(
        text: str, chunk_size: int = 400, overlap: int = 80
    ) -> List[str]:
        """Split text into overlapping chunks for embedding."""
        if not text or not text.strip():
            return []

        paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
        full_text = "\n".join(paragraphs)
        words = full_text.split()

        if len(words) <= chunk_size:
            return [full_text]

        chunks = []
        start = 0

        while start < len(words):
            end = min(start + chunk_size, len(words))
            chunk = " ".join(words[start:end])
            if chunk.strip():
                chunks.append(chunk.strip())

            if end >= len(words):
                break

            start += chunk_size - overlap

        return chunks

    @staticmethod
    def extract_key_info(text: str) -> dict:
        """Extract basic key information from document text."""
        info = {
            "has_email": False,
            "has_phone": False,
            "word_count": len(text.split()),
            "line_count": len(text.split("\n")),
        }

        import re

        if re.search(r"[\w.+-]+@[\w-]+\.[\w.-]+", text):
            info["has_email"] = True
        if re.search(r"[\+]?[\d\s\-\(\)]{7,15}", text):
            info["has_phone"] = True

        return info
