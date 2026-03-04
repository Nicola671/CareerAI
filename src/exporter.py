"""
Export Module - Generate PDF, DOCX, TXT, and HTML downloads from AI responses.
Premium formatting with professional layouts and smart content detection.
"""
import io
import re
from datetime import datetime
from typing import List, Dict, Optional


# ======================== TEXT CLEANING ========================

def clean_markdown(text: str) -> str:
    """Remove markdown formatting for clean document export."""
    # Remove bold/italic markers
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'__(.+?)__', r'\1', text)
    text = re.sub(r'_(.+?)_', r'\1', text)
    # Remove headers markers but keep text
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
    # Remove bullet markers
    text = re.sub(r'^[\-\*]\s+', '- ', text, flags=re.MULTILINE)
    # Remove numbered lists prefix (keep number)
    text = re.sub(r'^(\d+)\.\s+', r'\1. ', text, flags=re.MULTILINE)
    # Remove code blocks markers
    text = re.sub(r'```[\w]*\n?', '', text)
    # Remove inline code
    text = re.sub(r'`(.+?)`', r'\1', text)
    # Remove links but keep text
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    # Remove emojis (common ones)
    text = re.sub(
        r'[🎯✉️📈🧠💡🚀⭐📋🔍💼📊✅❌⚠️🎓🗺️📄🤖👋💬📎📷📚📭🎨🔥💪🏆🌟✨🎉💰📌🔑⚡🛠️🏁📐💎🥇🥈🥉]',
        '', text
    )
    # Clean up extra whitespace
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def detect_content_type(text: str) -> str:
    """Detect the type of content for smart file naming."""
    text_lower = text.lower()
    if any(w in text_lower for w in ['cover letter', 'carta de presentación', 'carta de motivación', 'estimado', 'dear']):
        return "cover_letter"
    if any(w in text_lower for w in ['match', 'compatibilidad', 'porcentaje', 'afinidad', '% de match']):
        return "job_match"
    if any(w in text_lower for w in ['skills gap', 'habilidades faltantes', 'roadmap', 'skill gap', 'brecha']):
        return "skills_analysis"
    if any(w in text_lower for w in ['resumen', 'perfil profesional', 'summary', 'strengths']):
        return "profile_summary"
    return "response"


def get_smart_filename(text: str, extension: str) -> str:
    """Generate an intelligent filename based on content type."""
    content_type = detect_content_type(text)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    
    type_names = {
        "cover_letter": "CoverLetter",
        "job_match": "JobMatch",
        "skills_analysis": "SkillsAnalysis",
        "profile_summary": "ProfileSummary",
        "response": "CareerAI",
    }
    
    name = type_names.get(content_type, "CareerAI")
    return f"{name}_{timestamp}.{extension}"


def get_smart_title(text: str) -> str:
    """Generate a smart document title based on content."""
    content_type = detect_content_type(text)
    titles = {
        "cover_letter": "Carta de Presentación",
        "job_match": "Análisis de Compatibilidad",
        "skills_analysis": "Análisis de Habilidades",
        "profile_summary": "Resumen de Perfil",
        "response": "Respuesta CareerAI",
    }
    return titles.get(content_type, "Respuesta CareerAI")


# ======================== MARKDOWN PARSER ========================

def parse_markdown_blocks(text: str) -> list:
    """
    Parse markdown into structured blocks for rich document rendering.
    Returns list of dicts: {type, content, level}
    """
    blocks = []
    lines = text.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Headers
        header_match = re.match(r'^(#{1,6})\s+(.+)', line)
        if header_match:
            level = len(header_match.group(1))
            blocks.append({
                'type': 'header',
                'content': header_match.group(2).strip(),
                'level': level
            })
            i += 1
            continue
        
        # Horizontal rules
        if re.match(r'^[\-\*\_]{3,}\s*$', line):
            blocks.append({'type': 'hr', 'content': '', 'level': 0})
            i += 1
            continue
        
        # Bullet lists
        bullet_match = re.match(r'^[\-\*]\s+(.+)', line)
        if bullet_match:
            items = [bullet_match.group(1).strip()]
            i += 1
            while i < len(lines):
                next_bullet = re.match(r'^[\-\*]\s+(.+)', lines[i])
                if next_bullet:
                    items.append(next_bullet.group(1).strip())
                    i += 1
                else:
                    break
            blocks.append({'type': 'bullet_list', 'content': items, 'level': 0})
            continue
        
        # Numbered lists
        num_match = re.match(r'^(\d+)\.\s+(.+)', line)
        if num_match:
            items = [num_match.group(2).strip()]
            i += 1
            while i < len(lines):
                next_num = re.match(r'^\d+\.\s+(.+)', lines[i])
                if next_num:
                    items.append(next_num.group(1).strip())
                    i += 1
                else:
                    break
            blocks.append({'type': 'numbered_list', 'content': items, 'level': 0})
            continue
        
        # Code blocks
        if line.strip().startswith('```'):
            lang = line.strip()[3:]
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1  # skip closing ```
            blocks.append({
                'type': 'code',
                'content': '\n'.join(code_lines),
                'level': 0,
                'lang': lang
            })
            continue
        
        # Bold/emphasis lines (like "**Sección:**")
        bold_match = re.match(r'^\*\*(.+?)\*\*:?\s*$', line.strip())
        if bold_match:
            blocks.append({
                'type': 'bold_heading',
                'content': bold_match.group(1).strip(),
                'level': 0
            })
            i += 1
            continue
        
        # Empty lines
        if not line.strip():
            i += 1
            continue
        
        # Regular paragraph (collect consecutive lines)
        para_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not re.match(r'^#{1,6}\s+', lines[i]) \
                and not re.match(r'^[\-\*]\s+', lines[i]) and not re.match(r'^\d+\.\s+', lines[i]) \
                and not lines[i].strip().startswith('```') and not re.match(r'^\*\*(.+?)\*\*:?\s*$', lines[i].strip()):
            para_lines.append(lines[i])
            i += 1
        
        blocks.append({
            'type': 'paragraph',
            'content': ' '.join(para_lines),
            'level': 0
        })
    
    return blocks


def strip_inline_md(text: str) -> str:
    """Remove inline markdown (bold, italic, code, links) from text."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'__(.+?)__', r'\1', text)
    text = re.sub(r'_(.+?)_', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    return text


def _sanitize_for_pdf(text: str) -> str:
    """Replace Unicode characters with ASCII equivalents for PDF Helvetica font."""
    replacements = {
        '\u2022': '-',
        '\u2013': '-',
        '\u2014': '--',
        '\u2018': "'",
        '\u2019': "'",
        '\u201c': '"',
        '\u201d': '"',
        '\u2026': '...',
        '\u2192': '->',
        '\u2190': '<-',
        '\u00b7': '-',
        '\u2500': '-',
        '\u2501': '-',
        '\u25cf': '-',
        '\u2605': '*',
        '\u2713': 'v',
        '\u2717': 'x',
    }
    for char, replacement in replacements.items():
        text = text.replace(char, replacement)
    text = re.sub(
        r'[\U0001F300-\U0001F9FF\U00002702-\U000027B0\U0000FE00-\U0000FE0F\U0001FA00-\U0001FA6F\U0001FA70-\U0001FAFF]',
        '', text
    )
    return text


# ======================== PDF EXPORT ========================

def export_to_pdf(text: str, title: Optional[str] = None) -> bytes:
    """Export text content to a premium-styled PDF."""
    try:
        from fpdf import FPDF
    except ImportError:
        raise ValueError("Instala fpdf2: pip install fpdf2")

    if title is None:
        title = get_smart_title(text)

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=25)
    pdf.add_page()
    
    page_width = pdf.w - 40  # margins

    # ---- Header Band ----
    pdf.set_fill_color(88, 60, 200)
    pdf.rect(0, 0, 210, 3, 'F')

    # ---- Title ----
    pdf.set_y(15)
    pdf.set_font("Helvetica", "B", 20)
    pdf.set_text_color(88, 60, 200)
    pdf.cell(0, 12, title, new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(1)

    # ---- Subtitle / Date ----
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(140, 140, 150)
    date_str = datetime.now().strftime("%d de %B, %Y | %H:%M")
    pdf.cell(0, 6, f"Generado el {date_str}", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(2)

    # ---- Divider ----
    y = pdf.get_y()
    pdf.set_draw_color(200, 200, 215)
    pdf.set_line_width(0.3)
    # Gradient-like effect with multiple lines
    pdf.set_draw_color(88, 60, 200)
    pdf.line(70, y, 140, y)
    pdf.set_draw_color(200, 200, 215)
    pdf.line(40, y + 0.5, 170, y + 0.5)
    pdf.ln(10)

    # ---- Content Blocks ----
    blocks = parse_markdown_blocks(text)

    for block in blocks:
        btype = block['type']

        if btype == 'header':
            level = block['level']
            content = _sanitize_for_pdf(strip_inline_md(block['content']))
            pdf.ln(4)
            
            if level == 1:
                pdf.set_font("Helvetica", "B", 16)
                pdf.set_text_color(30, 30, 45)
            elif level == 2:
                pdf.set_font("Helvetica", "B", 14)
                pdf.set_text_color(88, 60, 200)
            elif level == 3:
                pdf.set_font("Helvetica", "B", 12)
                pdf.set_text_color(60, 60, 80)
            else:
                pdf.set_font("Helvetica", "B", 11)
                pdf.set_text_color(80, 80, 100)
            
            pdf.multi_cell(0, 7, content)
            pdf.ln(2)

        elif btype == 'bold_heading':
            content = _sanitize_for_pdf(strip_inline_md(block['content']))
            pdf.ln(3)
            pdf.set_font("Helvetica", "B", 12)
            pdf.set_text_color(88, 60, 200)
            pdf.multi_cell(0, 7, content)
            pdf.set_text_color(30, 30, 45)
            pdf.ln(1)

        elif btype == 'paragraph':
            content = _sanitize_for_pdf(strip_inline_md(block['content']))
            pdf.set_font("Helvetica", "", 10.5)
            pdf.set_text_color(40, 40, 50)
            pdf.multi_cell(0, 5.5, content)
            pdf.ln(3)

        elif btype == 'bullet_list':
            for item in block['content']:
                item_clean = _sanitize_for_pdf(strip_inline_md(item))
                pdf.set_font("Helvetica", "", 10.5)
                pdf.set_text_color(88, 60, 200)
                pdf.cell(8, 5.5, "-")
                pdf.set_text_color(40, 40, 50)
                pdf.multi_cell(0, 5.5, f" {item_clean}")
                pdf.ln(1)
            pdf.ln(2)

        elif btype == 'numbered_list':
            for idx, item in enumerate(block['content'], 1):
                item_clean = _sanitize_for_pdf(strip_inline_md(item))
                pdf.set_font("Helvetica", "B", 10.5)
                pdf.set_text_color(88, 60, 200)
                pdf.cell(10, 5.5, f"{idx}.")
                pdf.set_font("Helvetica", "", 10.5)
                pdf.set_text_color(40, 40, 50)
                pdf.multi_cell(0, 5.5, f" {item_clean}")
                pdf.ln(1)
            pdf.ln(2)

        elif btype == 'code':
            pdf.ln(2)
            # Code block background
            pdf.set_fill_color(245, 245, 248)
            pdf.set_font("Courier", "", 9)
            pdf.set_text_color(60, 60, 80)
            code_lines = block['content'].split('\n')
            for cl in code_lines:
                pdf.cell(0, 5, f"  {cl}", new_x="LMARGIN", new_y="NEXT", fill=True)
            pdf.ln(3)

        elif btype == 'hr':
            pdf.ln(3)
            y = pdf.get_y()
            pdf.set_draw_color(200, 200, 215)
            pdf.line(20, y, 190, y)
            pdf.ln(5)

    # ---- Footer ----
    pdf.ln(10)
    y = pdf.get_y()
    pdf.set_draw_color(88, 60, 200)
    pdf.line(60, y, 150, y)
    pdf.ln(6)
    pdf.set_font("Helvetica", "I", 8)
    pdf.set_text_color(160, 160, 175)
    pdf.cell(0, 5, "Generado por CareerAI - Asistente de Carrera con IA", align="C")
    pdf.ln(4)
    pdf.set_font("Helvetica", "", 7)
    pdf.cell(0, 4, "Powered by RAG + Llama 3.3 + ChromaDB", align="C")

    # Bottom band
    pdf.set_fill_color(88, 60, 200)
    pdf.rect(0, 294, 210, 3, 'F')

    return pdf.output()


# ======================== DOCX EXPORT ========================

def export_to_docx(text: str, title: Optional[str] = None) -> bytes:
    """Export text content to a professionally styled DOCX."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise ValueError("Instala python-docx: pip install python-docx")

    if title is None:
        title = get_smart_title(text)

    doc = Document()

    # ---- Page margins ----
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ---- Default font ----
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(40, 40, 50)

    # ---- Accent line ----
    accent_para = doc.add_paragraph()
    accent_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    accent_run = accent_para.add_run("━" * 40)
    accent_run.font.color.rgb = RGBColor(88, 60, 200)
    accent_run.font.size = Pt(6)

    # ---- Title ----
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.space_after = Pt(4)
    title_run = title_para.add_run(title)
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(88, 60, 200)
    title_run.font.name = 'Calibri Light'

    # ---- Date ----
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.space_after = Pt(2)
    date_str = datetime.now().strftime("%d de %B, %Y • %H:%M")
    date_run = date_para.add_run(f"Generado el {date_str}")
    date_run.font.size = Pt(9)
    date_run.font.color.rgb = RGBColor(140, 140, 150)

    # ---- Divider ----
    div_para = doc.add_paragraph()
    div_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    div_para.space_after = Pt(12)
    div_run = div_para.add_run("─" * 50)
    div_run.font.color.rgb = RGBColor(200, 200, 215)
    div_run.font.size = Pt(8)

    # ---- Content Blocks ----
    blocks = parse_markdown_blocks(text)

    for block in blocks:
        btype = block['type']

        if btype == 'header':
            level = block['level']
            content = strip_inline_md(block['content'])
            
            p = doc.add_paragraph()
            p.space_before = Pt(12)
            p.space_after = Pt(4)
            run = p.add_run(content)
            run.font.bold = True
            
            if level == 1:
                run.font.size = Pt(18)
                run.font.color.rgb = RGBColor(30, 30, 45)
            elif level == 2:
                run.font.size = Pt(15)
                run.font.color.rgb = RGBColor(88, 60, 200)
            elif level == 3:
                run.font.size = Pt(13)
                run.font.color.rgb = RGBColor(60, 60, 80)
            else:
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(80, 80, 100)

        elif btype == 'bold_heading':
            content = strip_inline_md(block['content'])
            p = doc.add_paragraph()
            p.space_before = Pt(8)
            p.space_after = Pt(2)
            run = p.add_run(content)
            run.font.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(88, 60, 200)

        elif btype == 'paragraph':
            content = strip_inline_md(block['content'])
            p = doc.add_paragraph(content)
            p.paragraph_format.line_spacing = Pt(16)
            p.space_after = Pt(6)

        elif btype == 'bullet_list':
            for item in block['content']:
                item_clean = strip_inline_md(item)
                p = doc.add_paragraph(item_clean, style='List Bullet')
                p.paragraph_format.line_spacing = Pt(15)

        elif btype == 'numbered_list':
            for item in block['content']:
                item_clean = strip_inline_md(item)
                p = doc.add_paragraph(item_clean, style='List Number')
                p.paragraph_format.line_spacing = Pt(15)

        elif btype == 'code':
            code_para = doc.add_paragraph()
            code_para.space_before = Pt(6)
            code_para.space_after = Pt(6)
            # Add shading to code block
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'F5F5F8')
            shading.set(qn('w:val'), 'clear')
            code_para.paragraph_format.element.get_or_add_pPr().append(shading)
            
            run = code_para.add_run(block['content'])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(60, 60, 80)

        elif btype == 'hr':
            hr_para = doc.add_paragraph()
            hr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            hr_run = hr_para.add_run("─" * 50)
            hr_run.font.color.rgb = RGBColor(200, 200, 215)
            hr_run.font.size = Pt(8)

    # ---- Footer ----
    div_para2 = doc.add_paragraph()
    div_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    div_para2.space_before = Pt(20)
    div_run2 = div_para2.add_run("─" * 50)
    div_run2.font.color.rgb = RGBColor(200, 200, 215)
    div_run2.font.size = Pt(8)

    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        "Generado por CareerAI — Asistente de Carrera con IA"
    )
    footer_run.font.size = Pt(8)
    footer_run.font.italic = True
    footer_run.font.color.rgb = RGBColor(160, 160, 175)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run("Powered by RAG + Llama 3.3 + ChromaDB")
    sub_run.font.size = Pt(7)
    sub_run.font.color.rgb = RGBColor(180, 180, 195)

    # ---- Save to bytes ----
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ======================== TXT EXPORT ========================

def export_to_txt(text: str) -> bytes:
    """Export text content as a clean, well-formatted TXT file."""
    clean = clean_markdown(text)
    title = get_smart_title(text)
    date_str = datetime.now().strftime('%d/%m/%Y %H:%M')
    
    header = (
        f"{'=' * 60}\n"
        f"  {title}\n"
        f"  Generado: {date_str}\n"
        f"  CareerAI — Asistente de Carrera con IA\n"
        f"{'=' * 60}\n\n"
    )
    
    footer = (
        f"\n\n{'─' * 60}\n"
        f"Generado por CareerAI | Powered by RAG + Llama 3.3\n"
    )
    
    return (header + clean + footer).encode("utf-8")


# ======================== HTML EXPORT ========================

def export_to_html(text: str, title: Optional[str] = None) -> bytes:
    """Export text content as a beautifully styled standalone HTML file."""
    import html as html_lib

    if title is None:
        title = get_smart_title(text)

    date_str = datetime.now().strftime("%d de %B, %Y • %H:%M")
    
    # Convert markdown to HTML-like content
    blocks = parse_markdown_blocks(text)
    content_html = ""
    
    for block in blocks:
        btype = block['type']
        
        if btype == 'header':
            level = block['level']
            content = html_lib.escape(strip_inline_md(block['content']))
            tag = f"h{min(level + 1, 6)}"  # shift down since h1 is title
            content_html += f"<{tag}>{content}</{tag}>\n"
        
        elif btype == 'bold_heading':
            content = html_lib.escape(strip_inline_md(block['content']))
            content_html += f'<h3 class="accent">{content}</h3>\n'
        
        elif btype == 'paragraph':
            content = html_lib.escape(strip_inline_md(block['content']))
            content_html += f"<p>{content}</p>\n"
        
        elif btype == 'bullet_list':
            content_html += "<ul>\n"
            for item in block['content']:
                item_clean = html_lib.escape(strip_inline_md(item))
                content_html += f"  <li>{item_clean}</li>\n"
            content_html += "</ul>\n"
        
        elif btype == 'numbered_list':
            content_html += "<ol>\n"
            for item in block['content']:
                item_clean = html_lib.escape(strip_inline_md(item))
                content_html += f"  <li>{item_clean}</li>\n"
            content_html += "</ol>\n"
        
        elif btype == 'code':
            lang = block.get('lang', '')
            code_content = html_lib.escape(block['content'])
            content_html += f'<pre><code class="{lang}">{code_content}</code></pre>\n'
        
        elif btype == 'hr':
            content_html += '<hr>\n'

    html_template = f"""<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{html_lib.escape(title)} — CareerAI</title>
    <link href="https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap" rel="stylesheet">
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        
        body {{
            font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
            background: linear-gradient(135deg, #0f0f23 0%, #1a1a2e 30%, #16213e 60%, #0f0f23 100%);
            color: #e4e4e7;
            min-height: 100vh;
            line-height: 1.7;
        }}
        
        .container {{
            max-width: 780px;
            margin: 0 auto;
            padding: 40px 30px;
        }}
        
        .header {{
            text-align: center;
            margin-bottom: 40px;
            padding-bottom: 30px;
            border-bottom: 1px solid rgba(139, 92, 246, 0.2);
            position: relative;
        }}
        
        .header::before {{
            content: '';
            position: absolute;
            bottom: -1px;
            left: 50%;
            transform: translateX(-50%);
            width: 120px;
            height: 2px;
            background: linear-gradient(90deg, transparent, #8b5cf6, transparent);
        }}
        
        .brand {{
            font-size: 0.75rem;
            font-weight: 600;
            text-transform: uppercase;
            letter-spacing: 0.15em;
            color: #8b5cf6;
            margin-bottom: 12px;
        }}
        
        h1 {{
            font-size: 2rem;
            font-weight: 700;
            background: linear-gradient(135deg, #a78bfa, #c084fc, #e879f9, #f472b6);
            background-clip: text;
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            margin-bottom: 8px;
            letter-spacing: -0.02em;
        }}
        
        .date {{
            font-size: 0.85rem;
            color: #71717a;
        }}
        
        .content {{
            background: rgba(24, 24, 27, 0.5);
            border: 1px solid rgba(63, 63, 70, 0.3);
            border-radius: 20px;
            padding: 40px 36px;
            backdrop-filter: blur(20px);
            box-shadow: 0 25px 60px -15px rgba(0, 0, 0, 0.4);
        }}
        
        h2 {{
            font-size: 1.5rem;
            font-weight: 700;
            color: #fafafa;
            margin: 28px 0 12px 0;
            letter-spacing: -0.01em;
        }}
        
        h3 {{
            font-size: 1.2rem;
            font-weight: 600;
            color: #d4d4d8;
            margin: 22px 0 10px 0;
        }}
        
        h3.accent {{
            color: #a78bfa;
        }}
        
        h4, h5, h6 {{
            font-size: 1rem;
            font-weight: 600;
            color: #a1a1aa;
            margin: 18px 0 8px 0;
        }}
        
        p {{
            margin: 0 0 14px 0;
            color: #d4d4d8;
            font-size: 0.95rem;
        }}
        
        ul, ol {{
            margin: 10px 0 18px 0;
            padding-left: 24px;
        }}
        
        li {{
            margin-bottom: 8px;
            color: #d4d4d8;
            font-size: 0.95rem;
        }}
        
        li::marker {{
            color: #8b5cf6;
        }}
        
        pre {{
            background: rgba(15, 15, 30, 0.6);
            border: 1px solid rgba(63, 63, 70, 0.3);
            border-radius: 12px;
            padding: 18px 20px;
            overflow-x: auto;
            margin: 14px 0;
        }}
        
        code {{
            font-family: 'JetBrains Mono', 'Fira Code', Consolas, monospace;
            font-size: 0.85rem;
            color: #c4b5fd;
        }}
        
        hr {{
            border: none;
            height: 1px;
            background: linear-gradient(90deg, transparent, rgba(139, 92, 246, 0.3), transparent);
            margin: 24px 0;
        }}
        
        .footer {{
            text-align: center;
            margin-top: 40px;
            padding-top: 24px;
            border-top: 1px solid rgba(63, 63, 70, 0.2);
        }}
        
        .footer p {{
            color: #52525b;
            font-size: 0.78rem;
        }}
        
        .footer .powered {{
            font-size: 0.72rem;
            color: #3f3f46;
            margin-top: 4px;
        }}
        
        @media print {{
            body {{ background: white; color: #1a1a2e; }}
            .content {{ border: 1px solid #e5e7eb; box-shadow: none; background: white; }}
            h1 {{ color: #4c1d95; -webkit-text-fill-color: #4c1d95; }}
            h2 {{ color: #1a1a2e; }}
            h3, h3.accent {{ color: #4c1d95; }}
            p, li {{ color: #374151; }}
            pre {{ background: #f9fafb; border: 1px solid #e5e7eb; }}
            code {{ color: #6d28d9; }}
        }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <div class="brand">CareerAI</div>
            <h1>{html_lib.escape(title)}</h1>
            <div class="date">{date_str}</div>
        </div>
        
        <div class="content">
            {content_html}
        </div>
        
        <div class="footer">
            <p>Generado por CareerAI — Asistente de Carrera con IA</p>
            <p class="powered">Powered by RAG + Llama 3.3 + ChromaDB</p>
        </div>
    </div>
</body>
</html>"""

    return html_template.encode("utf-8")


# ======================== CONVERSATION EXPORT ========================

def export_conversation_to_pdf(messages: List[Dict], title: str = "Conversación CareerAI") -> bytes:
    """Export full conversation history to PDF."""
    try:
        from fpdf import FPDF
    except ImportError:
        raise ValueError("Instala fpdf2: pip install fpdf2")

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    # Header band
    pdf.set_fill_color(88, 60, 200)
    pdf.rect(0, 0, 210, 3, 'F')

    # Title
    pdf.set_y(15)
    pdf.set_font("Helvetica", "B", 18)
    pdf.set_text_color(88, 60, 200)
    pdf.cell(0, 12, title, new_x="LMARGIN", new_y="NEXT", align="C")

    # Date
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(140, 140, 150)
    date_str = datetime.now().strftime("%d/%m/%Y %H:%M")
    pdf.cell(0, 6, f"Exportado el {date_str}", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(4)

    # Stats
    user_msgs = sum(1 for m in messages if m["role"] == "user")
    ai_msgs = sum(1 for m in messages if m["role"] == "assistant")
    pdf.set_font("Helvetica", "", 8)
    pdf.set_text_color(160, 160, 175)
    pdf.cell(0, 5, f"{user_msgs} preguntas · {ai_msgs} respuestas · {len(messages)} mensajes totales",
             new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(6)

    # Divider
    y = pdf.get_y()
    pdf.set_draw_color(200, 200, 215)
    pdf.line(20, y, 190, y)
    pdf.ln(8)

    # Messages
    for i, msg in enumerate(messages):
        is_user = msg["role"] == "user"
        
        # Role label
        pdf.set_font("Helvetica", "B", 10)
        if is_user:
            pdf.set_text_color(100, 100, 120)
            pdf.cell(0, 6, f"Tu ({i + 1})", new_x="LMARGIN", new_y="NEXT")
        else:
            pdf.set_text_color(88, 60, 200)
            pdf.cell(0, 6, f"CareerAI ({i + 1})", new_x="LMARGIN", new_y="NEXT")
        
        # Content
        clean = _sanitize_for_pdf(clean_markdown(msg["content"]))
        pdf.set_font("Helvetica", "", 10)
        pdf.set_text_color(40, 40, 50)
        
        for paragraph in clean.split('\n'):
            paragraph = paragraph.strip()
            if not paragraph:
                pdf.ln(2)
                continue
            if paragraph.startswith('•'):
                pdf.multi_cell(0, 5, paragraph)
                pdf.ln(1)
            else:
                pdf.multi_cell(0, 5, paragraph)
                pdf.ln(1)
        
        pdf.ln(4)
        
        # Separator between messages
        if i < len(messages) - 1:
            y = pdf.get_y()
            pdf.set_draw_color(220, 220, 230)
            pdf.set_line_width(0.2)
            pdf.line(30, y, 180, y)
            pdf.ln(5)

    # Footer
    pdf.ln(8)
    pdf.set_draw_color(88, 60, 200)
    pdf.line(60, pdf.get_y(), 150, pdf.get_y())
    pdf.ln(6)
    pdf.set_font("Helvetica", "I", 8)
    pdf.set_text_color(160, 160, 175)
    pdf.cell(0, 5, "CareerAI - Asistente de Carrera con IA", align="C")

    # Bottom band
    pdf.set_fill_color(88, 60, 200)
    pdf.rect(0, 294, 210, 3, 'F')

    return pdf.output()


def export_conversation_to_docx(messages: List[Dict], title: str = "Conversación CareerAI") -> bytes:
    """Export full conversation history to DOCX (Word)."""
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ValueError("Instala python-docx: pip install python-docx")

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(40, 40, 50)

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.space_after = Pt(4)
    title_run = title_para.add_run(title)
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(88, 60, 200)

    # Date & stats
    date_str = datetime.now().strftime("%d de %B, %Y • %H:%M")
    user_msgs = sum(1 for m in messages if m["role"] == "user")
    ai_msgs = sum(1 for m in messages if m["role"] == "assistant")
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.space_after = Pt(2)
    date_run = date_para.add_run(f"Exportado el {date_str}")
    date_run.font.size = Pt(9)
    date_run.font.color.rgb = RGBColor(140, 140, 150)

    stats_para = doc.add_paragraph()
    stats_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    stats_para.space_after = Pt(12)
    stats_run = stats_para.add_run(f"{user_msgs} preguntas · {ai_msgs} respuestas · {len(messages)} mensajes")
    stats_run.font.size = Pt(8)
    stats_run.font.color.rgb = RGBColor(160, 160, 175)

    # Divider
    div_para = doc.add_paragraph()
    div_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    div_para.space_after = Pt(12)
    div_run = div_para.add_run("─" * 50)
    div_run.font.color.rgb = RGBColor(200, 200, 215)
    div_run.font.size = Pt(8)

    # Messages
    for i, msg in enumerate(messages):
        is_user = msg["role"] == "user"
        role_label = f"Tú (#{i + 1})" if is_user else f"CareerAI (#{i + 1})"

        role_para = doc.add_paragraph()
        role_para.space_before = Pt(14)
        role_para.space_after = Pt(4)
        role_run = role_para.add_run(role_label)
        role_run.font.bold = True
        role_run.font.size = Pt(11)
        if is_user:
            role_run.font.color.rgb = RGBColor(80, 80, 100)
        else:
            role_run.font.color.rgb = RGBColor(88, 60, 200)

        clean = clean_markdown(msg["content"])
        for line in clean.split("\n"):
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue
            p = doc.add_paragraph(line)
            p.paragraph_format.line_spacing = Pt(15)
            p.paragraph_format.space_after = Pt(4)

        if i < len(messages) - 1:
            sep = doc.add_paragraph()
            sep.space_after = Pt(6)
            sep_run = sep.add_run("─" * 40)
            sep_run.font.color.rgb = RGBColor(220, 220, 230)
            sep_run.font.size = Pt(6)

    # Footer
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run("CareerAI — Asistente de Carrera con IA")
    footer_run.font.size = Pt(8)
    footer_run.font.italic = True
    footer_run.font.color.rgb = RGBColor(160, 160, 175)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def export_conversation_to_html(messages: List[Dict], title: str = "Conversación CareerAI") -> bytes:
    """Export full conversation as a beautifully styled HTML file."""
    import html as html_lib

    date_str = datetime.now().strftime("%d de %B, %Y • %H:%M")
    user_msgs = sum(1 for m in messages if m["role"] == "user")
    ai_msgs = sum(1 for m in messages if m["role"] == "assistant")

    messages_html = ""
    for i, msg in enumerate(messages):
        is_user = msg["role"] == "user"
        role_class = "user-msg" if is_user else "ai-msg"
        role_label = "Tú" if is_user else "CareerAI"
        avatar = "👤" if is_user else "🤖"
        clean = html_lib.escape(clean_markdown(msg["content"]))
        # Convert newlines to <br>
        clean = clean.replace('\n', '<br>')
        
        messages_html += f"""
        <div class="message {role_class}">
            <div class="message-header">
                <span class="avatar">{avatar}</span>
                <span class="role">{role_label}</span>
                <span class="msg-num">#{i + 1}</span>
            </div>
            <div class="message-body">{clean}</div>
        </div>
        """

    html_content = f"""<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{html_lib.escape(title)} — CareerAI</title>
    <link href="https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap" rel="stylesheet">
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{
            font-family: 'Inter', sans-serif;
            background: linear-gradient(135deg, #0f0f23, #1a1a2e, #16213e, #0f0f23);
            color: #e4e4e7;
            min-height: 100vh;
            line-height: 1.6;
        }}
        .container {{ max-width: 800px; margin: 0 auto; padding: 40px 24px; }}
        .header {{
            text-align: center;
            margin-bottom: 32px;
            padding-bottom: 24px;
            border-bottom: 1px solid rgba(139, 92, 246, 0.2);
        }}
        .brand {{ font-size: 0.75rem; font-weight: 600; text-transform: uppercase; letter-spacing: 0.15em; color: #8b5cf6; margin-bottom: 10px; }}
        h1 {{
            font-size: 1.8rem; font-weight: 700;
            background: linear-gradient(135deg, #a78bfa, #c084fc, #e879f9);
            -webkit-background-clip: text; -webkit-text-fill-color: transparent;
            margin-bottom: 6px;
        }}
        .meta {{ color: #71717a; font-size: 0.85rem; }}
        .stats {{ color: #52525b; font-size: 0.8rem; margin-top: 6px; }}
        
        .message {{
            margin-bottom: 16px;
            border-radius: 16px;
            padding: 18px 22px;
            border: 1px solid rgba(63, 63, 70, 0.3);
        }}
        .user-msg {{
            background: rgba(24, 24, 27, 0.4);
        }}
        .ai-msg {{
            background: rgba(88, 60, 200, 0.06);
            border-color: rgba(139, 92, 246, 0.15);
        }}
        .message-header {{
            display: flex;
            align-items: center;
            gap: 8px;
            margin-bottom: 10px;
        }}
        .avatar {{ font-size: 1.2rem; }}
        .role {{ font-weight: 600; font-size: 0.85rem; color: #a1a1aa; }}
        .ai-msg .role {{ color: #a78bfa; }}
        .msg-num {{ font-size: 0.72rem; color: #52525b; margin-left: auto; }}
        .message-body {{ font-size: 0.92rem; color: #d4d4d8; line-height: 1.7; }}
        
        .footer {{
            text-align: center;
            margin-top: 32px;
            padding-top: 20px;
            border-top: 1px solid rgba(63, 63, 70, 0.2);
        }}
        .footer p {{ color: #52525b; font-size: 0.78rem; }}
        
        @media print {{
            body {{ background: white; color: #1a1a2e; }}
            .message {{ border: 1px solid #e5e7eb; }}
            .ai-msg {{ background: #f8f5ff; }}
            .message-body {{ color: #374151; }}
        }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <div class="brand">CareerAI</div>
            <h1>{html_lib.escape(title)}</h1>
            <div class="meta">{date_str}</div>
            <div class="stats">{user_msgs} preguntas · {ai_msgs} respuestas</div>
        </div>
        
        {messages_html}
        
        <div class="footer">
            <p>Generado por CareerAI — Asistente de Carrera con IA</p>
        </div>
    </div>
</body>
</html>"""

    return html_content.encode("utf-8")
